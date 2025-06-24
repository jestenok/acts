import datetime
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from settings import ORG_CREDS, MONTH_NUMBER_MAP, EMP_CREDS


def set_default_style(document):
    for section in document.sections:
        section.top_margin = Pt(30)
        section.left_margin = Pt(50)
        section.right_margin = Pt(30)
        section.bottom_margin = Pt(20)

    default_style = document.styles['Normal']
    default_style.font.name = 'Times New Roman'
    default_style.font.size = Pt(12)
    default_style.paragraph_format.line_spacing = Pt(12)
    default_style.paragraph_format.space_after = Pt(0)


def create_documents(org: str, tasks: dict, dones: dict, month: int, emp: str):
    base_dir = os.path.dirname(os.path.abspath(__file__))
    save_folder = os.path.join(base_dir, emp, 'data')

    creds = ORG_CREDS[emp][org]
    month = datetime.date(day=1, month=month, year=2025).strftime("%m")
    emp_creds = EMP_CREDS[emp]

    create_act(creds, org, dones, month, emp_creds, save_folder)
    create_tz(creds, org, tasks, month, emp_creds, save_folder)


def create_act(creds: dict, org: str, tasks: dict, month: str, emp_creds: dict, save_folder: str):
    d = Document()
    set_default_style(d)

    p = d.add_paragraph('АКТ №б/н')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = True

    p = d.add_paragraph('Оказанных услуг')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = True

    _ = d.add_paragraph('г. Москва' + ' ' * 125 + f'«30» {MONTH_NUMBER_MAP[month]} 2025 г.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _ = d.add_paragraph()
    p = d.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    r = p.add_run(f'''{creds['full_name']} ''')
    r.bold = True

    _ = p.add_run(
        f''' именуемое в дальнейшем «Заказчик», в {creds['directors_position_whos']} {creds['director_whos']}, действующего на основании Устава и Договора, с одной стороны и ''')

    r = p.add_run(f'{emp_creds["org_full"]}, ')
    r.bold = True

    _ = p.add_run(
        f'''именуемый в дальнейшем «Исполнитель», с другой стороны, вместе именуемые «Стороны», составили настоящий акт (далее - Акт) к договору оказания услуг №{creds['contract_number']} от {creds['contract_date']}, техническое задание №б/н от 01.{month}.2025 г.:)''')

    _ = d.add_paragraph(f'''
        1.Исполнителем оказаны следующие услуги:
    ''')

    for k, v in tasks.items():
        _ = d.add_paragraph(f'    {k}')
        for i in v:
            _ = d.add_paragraph(f'        - {i}')

    _ = d.add_paragraph(f'''
        2. Вознаграждение Исполнителя составляет ____,__ (_____ рублей, __ копеек)

        3. Стороны не имеют претензий друг к другу.

      4. Акт составлен и подписан в двух экземплярах, имеющих одинаковую юридическую силу, по одному для каждой из Сторон.

        5. Подписи сторон:
        ''')

    t = d.add_table(rows=3, cols=2)
    t.style = 'Table Grid'

    p = t.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Заказчик')
    r.bold = True

    p = t.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Исполнитель')
    r.bold = True

    c = t.cell(1, 0)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _ = p.add_run(creds['name'])
    p = c.add_paragraph(f'ИНН {creds["inn"]}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    c = t.cell(1, 1)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _ = p.add_run(f'{emp_creds["org"]}')
    p = c.add_paragraph(f'ИНН {emp_creds["inn"]}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    c = t.cell(2, 0)
    c.add_paragraph()
    short_dir = creds['director'].split(' ')
    short_dir = f'{short_dir[0]} {short_dir[1][0]}.{short_dir[2][0]}.'
    c.add_paragraph(f'________________ {short_dir}')

    c = t.cell(2, 1)
    c.add_paragraph()
    _ = c.add_paragraph(f'________________ {emp_creds["name"]}')

    d.save(f'''{save_folder}/01.{month}.2025 {org} акт.docx''')


def create_tz(creds: dict, org: str, todo: dict, month: str, emp_creds: dict, save_folder: str):
    d = Document()
    set_default_style(d)

    p = d.add_paragraph(
        '''
        Приложение № 1
        к договору №б/н
        от «04» апреля 2022 г.
        '''
    )
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p = d.add_paragraph(f'Техническое задание №б/н от 01.{month}.2025 г.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = True

    p = d.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(20)
    r = p.add_run(f'''{creds['full_name']}, ''')
    r.bold = True

    _ = p.add_run(
        f'''именуемое в дальнейшем «Заказчик», в {creds['directors_position_whos']} {creds['director_whos']}, действующего на основании Устава и Договора, с одной стороны и ''')

    r = p.add_run(f'{emp_creds["org_full"]}, ')
    r.bold = True

    _ = p.add_run('''именуемый в дальнейшем «Исполнитель», с другой стороны, вместе именуемые «Стороны», утвердили техническое задание в следующей редакции:
        ''')

    p = d.add_paragraph('Исполнитель обязуется оказать следующие услуги/работы:')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = True

    _ = d.add_paragraph('Задание:')

    for k, v in todo.items():
        _ = d.add_paragraph(k)
        for i in v:
            _ = d.add_paragraph(f'    - {i}')

    _ = d.add_paragraph('\nИТОГО:  стоимость услуг составит  ____,__ (_____ рублей __ копеек)\n')

    p = d.add_paragraph('Подписи сторон:')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = True

    t = d.add_table(rows=3, cols=2)
    t.style = 'Table Grid'

    p = t.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Заказчик')
    r.bold = True

    p = t.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Исполнитель')
    r.bold = True

    c = t.cell(1, 0)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _ = p.add_run(creds['name'])
    p = c.add_paragraph(f'ИНН {creds["inn"]}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    c = t.cell(1, 1)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _ = p.add_run(f'{emp_creds["org"]}')
    p = c.add_paragraph(f'ИНН {emp_creds["inn"]}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    c = t.cell(2, 0)
    c.add_paragraph()
    short_dir = creds['director'].split(' ')
    short_dir = f'{short_dir[0]} {short_dir[1][0]}.{short_dir[2][0]}.'
    c.add_paragraph(f'________________ {short_dir}')

    c = t.cell(2, 1)
    c.add_paragraph()
    _ = c.add_paragraph(f'________________ {emp_creds["name"]}')

    d.save(f'''{save_folder}/01.{month}.2025 {org} тз.docx''')
